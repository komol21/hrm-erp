"""
Service for extracting, validating, and sanitizing text from candidate CV files (PDF and DOCX).
"""

import os
import re
import io
from pypdf import PdfReader
from docx import Document


MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5MB
ALLOWED_EXTENSIONS = ['.pdf', '.docx']
MAX_CV_TEXT_LENGTH = 20000  # Character limit to optimize API token usage


class CVExtractionError(Exception):
    """Custom exception raised when CV file extraction or validation fails."""
    pass


def validate_cv_file(uploaded_file):
    """
    Validate uploaded file extension, size, and integrity.
    Raises CVExtractionError if validation fails.
    """
    if not uploaded_file:
        raise CVExtractionError("No file was uploaded.")

    # 1. Size check
    if uploaded_file.size > MAX_FILE_SIZE_BYTES:
        max_mb = MAX_FILE_SIZE_BYTES / (1024 * 1024)
        raise CVExtractionError(f"File size exceeds the {max_mb:.0f}MB limit. Uploaded file is {uploaded_file.size / (1024 * 1024):.2f}MB.")

    # 2. Extension check
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        allowed_str = ", ".join(ALLOWED_EXTENSIONS)
        raise CVExtractionError(f"Unsupported file format '{ext}'. Only {allowed_str} files are accepted.")

    return ext


def sanitize_text(text):
    """
    Sanitize extracted CV text:
    - Remove non-printable / control characters (except common whitespace)
    - Normalize repeated newlines and spaces
    - Truncate to maximum character limit
    """
    if not text:
        return ""

    # Replace null bytes and non-printable control characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

    # Normalize whitespace while preserving line structure
    lines = [line.strip() for line in text.splitlines()]
    # Remove excessive blank lines
    clean_lines = []
    prev_blank = False
    for line in lines:
        if line:
            clean_lines.append(line)
            prev_blank = False
        elif not prev_blank:
            clean_lines.append("")
            prev_blank = True

    cleaned = "\n".join(clean_lines).strip()
    return cleaned[:MAX_CV_TEXT_LENGTH]


def extract_text_from_pdf(file_obj):
    """Extract text content from a PDF file."""
    try:
        if hasattr(file_obj, 'read'):
            file_obj.seek(0)
            reader = PdfReader(file_obj)
        else:
            reader = PdfReader(file_obj)

        extracted_pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                extracted_pages.append(page_text)

        raw_text = "\n\n".join(extracted_pages)
        if not raw_text.strip():
            raise CVExtractionError("The uploaded PDF appears to be empty or contains only scanned images without selectable text.")

        return sanitize_text(raw_text)
    except CVExtractionError:
        raise
    except Exception as e:
        raise CVExtractionError(f"Failed to read PDF file: {str(e)}")


def extract_text_from_docx(file_obj):
    """Extract text content from a DOCX file."""
    try:
        if hasattr(file_obj, 'read'):
            file_obj.seek(0)
            # Use io.BytesIO if file_obj is in-memory
            content = file_obj.read()
            doc = Document(io.BytesIO(content))
        else:
            doc = Document(file_obj)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        all_text = "\n".join(paragraphs + table_texts)
        if not all_text.strip():
            raise CVExtractionError("The uploaded DOCX file appears to be empty.")

        return sanitize_text(all_text)
    except CVExtractionError:
        raise
    except Exception as e:
        raise CVExtractionError(f"Failed to read DOCX file: {str(e)}")


def extract_cv_text(uploaded_file):
    """
    Main entrypoint: validates uploaded file and extracts sanitized text.
    Returns: sanitized text (str).
    """
    ext = validate_cv_file(uploaded_file)

    if ext == '.pdf':
        return extract_text_from_pdf(uploaded_file)
    elif ext == '.docx':
        return extract_text_from_docx(uploaded_file)
    else:
        raise CVExtractionError(f"Unsupported file format: {ext}")


def extract_candidate_name_from_filename(filename):
    """
    Derive a natural candidate name from an uploaded CV filename.
    e.g. 'John_Doe_Resume_2026.pdf' -> 'John Doe'
         'alex-morgan-cv.docx' -> 'Alex Morgan'
    """
    if not filename:
        return 'Applicant'
    base = os.path.splitext(os.path.basename(filename))[0]
    tokens = re.split(r'[-_\s.]+', base)
    ignored = {'resume', 'cv', 'biodata', 'profile', 'updated', 'final', 'latest', 'v1', 'v2', 'new', 'draft', 'doc', 'docx', 'pdf'}
    name_tokens = [t for t in tokens if t.lower() not in ignored and not t.isdigit() and len(t) > 1]
    if name_tokens:
        return ' '.join(t.capitalize() for t in name_tokens)
    return base.replace('_', ' ').replace('-', ' ').title() or 'Applicant'


def extract_candidate_email_from_text(text):
    """Search for the first valid candidate email address inside extracted CV text."""
    if not text:
        return ''
    match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', text)
    return match.group(0).lower() if match else ''
